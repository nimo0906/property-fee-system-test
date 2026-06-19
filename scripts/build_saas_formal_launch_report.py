#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build formal DOCX/PDF launch report for SaaS commercial release.

Run with workspace Python that includes python-docx and reportlab.
"""

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / 'docs/saas-commercial-launch-report.md'
DOCX_OUT = ROOT / 'release/saas-commercial-launch-report.docx'
PDF_OUT = ROOT / 'release/saas-commercial-launch-report.pdf'
FONT = '宋体'

SECTIONS = [
    ('总体结论', [
        'SaaS 云端商业版员工后台主线已形成可部署、可演示、可验收、可签收的第一版商业交付闭环。',
        '当前版本面向通用 Linux/VPS、腾讯云、阿里云部署，覆盖登录、租户、权限、收费对象、收费项目、出账、审核、收款、报表、导入、导出、审计、备份和恢复演练。',
        '本阶段不包含业主端 H5、微信/支付宝真实支付、电子票据平台、独立授权云服务后台。',
    ]),
    ('P0-1 到 P0-13 完成状态', [
        'P0-1 业主与收费对象：业主档案、收费对象、导入映射、多租户隔离已完成。',
        'P0-2 计费规则：面积计费、固定金额、独立单价、服务期折算已完成。',
        'P0-3 批量出账：当前租户项目批量出账，重复账单跳过已完成。',
        'P0-4 收款欠费：部分收款、已收、欠费、报表联动已完成。',
        'P0-5 收据导出：账单/收款导出和收据详情含业务字段，不含内部字段。',
        'P0-6 导入复核：预览不写库、确认后写库、重复对象跳过已完成。',
        'P0-7 高风险审计：风险分级、详情脱敏、按租户隔离查询已完成。',
        'P0-8 备份恢复：备份记录、恢复演练记录、审计链路已完成。',
        'P0-9 商业上线验收总览：后台验收页和商业 readiness 检查已完成。',
        'P0-10 云端部署演练：通用 Linux/VPS、腾讯云、阿里云部署演练文档和检查脚本已完成。',
        'P0-11 商业交付演示：新租户从空库开始完成项目、收费对象、收费项目、出账、审核、收款、报表、导出、备份恢复演练。',
        'P0-12 最小上线包：部署配置、验收演示、运维备份、上线证据清单已完成。',
        'P0-13 客户交付签收：上线前人工验收和客户交付签收清单已完成。',
    ]),
    ('验证命令', [
        'python3 scripts/saas_minimal_launch_package_check.py',
        'python3 scripts/saas_customer_acceptance_signoff_check.py',
        'python3 scripts/saas_commercial_launch_report_check.py',
        'python3 scripts/saas_formal_launch_report_check.py',
        'python3 scripts/saas_release_gate.py',
    ]),
    ('部署资产', [
        'docker-compose.yml、.env.example、deploy/nginx/property-saas.conf、deploy/systemd/property-saas.service、deploy/logrotate/property-saas。',
        'docs/saas-cloud-ops-runbook.md、docs/saas-cloud-deployment-drill.md、docs/saas-minimal-launch-package.md。',
    ]),
    ('演示路径', [
        'scripts/saas_commercial_delivery_drill.py、scripts/saas_commercial_delivery_drill_check.py。',
        '后台商业验收页：/backoffice/acceptance；后台部署清单页：/backoffice/deploy-checklist。',
    ]),
    ('签收路径', [
        'docs/saas-customer-acceptance-signoff.md、release/saas-release-evidence.md、release/saas-isolation-evidence.md。',
    ]),
    ('数据和权限边界', [
        '租户数据隔离：不同公司之间收费对象、账单、收款、报表、审计、备份记录必须隔离。',
        '客户上传数据与系统自身数据隔离：客户文件、系统文件、备份和日志分目录保存。',
        '平台管理员可跨租户管理但必须审计；租户管理员只能管理本租户员工账号。',
    ]),
    ('上线证据', [
        '上线证据以 scripts/saas_release_gate.py 输出为准，并保留 release/saas-release-evidence.md、release/saas-isolation-evidence.md。',
    ]),
    ('后续建议', [
        '第一版上线后建议优先做真实客户小规模试点，先验证员工后台全量业务，再推进业主端 H5 和真实支付。',
        '正式接入微信/支付宝前，必须完成支付订单、回调签名、幂等、退款、对账和 HTTPS 公网域名部署。',
        '独立授权云服务后台后置，不与当前 SaaS 业务主库混在一起。',
    ]),
]

TABLES = [
    ('关键资产清单', [
        ['类别', '资产'],
        ['上线总报告', 'docs/saas-commercial-launch-report.md / release/saas-commercial-launch-report.docx / release/saas-commercial-launch-report.pdf'],
        ['客户签收', 'docs/saas-customer-acceptance-signoff.md'],
        ['最小上线包', 'docs/saas-minimal-launch-package.md'],
        ['商业演示', 'scripts/saas_commercial_delivery_drill.py'],
        ['总门禁', 'scripts/saas_release_gate.py'],
    ]),
    ('上线前必须确认', [
        ['确认项', '口径'],
        ['部署完成', 'Nginx HTTPS、app、PostgreSQL、systemd、logrotate 均完成。'],
        ['演示通过', '新租户空库业务闭环通过。'],
        ['恢复演练', '备份、metadata、checksum、恢复范围和业务核对通过。'],
        ['隔离通过', '租户数据隔离，客户上传数据与系统自身数据隔离。'],
    ]),
    ('客户签收核对表', [
        ['签收项', '验收口径'],
        ['部署完成确认', '服务、反向代理、HTTPS、日志轮转和运维脚本全部完成。'],
        ['最小上线包确认', '部署配置、运行服务、数据隔离、验收演示、运维备份、上线证据齐全。'],
        ['商业交付演示通过', '新租户从空库开始完成完整收费业务闭环。'],
        ['备份恢复演练通过', '备份文件、metadata、checksum 和恢复范围核对通过。'],
        ['租户隔离通过', '不同公司之间收费对象、账单、收款、报表、审计、备份记录隔离。'],
        ['账号权限确认', '平台管理员、租户管理员、财务、收费员、管理层边界清晰。'],
        ['数据不混用确认', '客户上传数据不混入系统自身目录，系统配置不混入客户上传目录。'],
        ['上线证据确认', '上线证据报告、租户隔离证据、签收清单均已留档。'],
    ]),
]


def set_run_font(run, size=12, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_format(p, first_indent=True):
    fmt = p.paragraph_format
    fmt.line_spacing = 0.8
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    if first_indent:
        fmt.first_line_indent = Pt(24)


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(p, first_indent=False)
    r = p.add_run(text)
    set_run_font(r, size=11, bold=bold)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_text(cell, cell.text, bold=(row_idx == 0))
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'E8EEF5' if row_idx == 0 else 'FFFFFF')
            tc_pr.append(shd)


def add_table(doc, title, rows):
    h = doc.add_paragraph()
    h.style = doc.styles['Heading 2']
    h.add_run(title)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    style_table(table)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    normal.font.size = Pt(12)
    for name, size in [('Heading 1', 14), ('Heading 2', 12), ('Heading 3', 12)]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 0.8
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)


def build_docx():
    doc = Document()
    configure_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 0.8
    r = title.add_run('SaaS 云端商业版上线总报告')
    set_run_font(r, size=22, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.line_spacing = 0.8
    set_run_font(subtitle.add_run('正式 Word/PDF 交付版'), size=12, bold=False)
    for heading, paras in SECTIONS:
        h = doc.add_paragraph(style='Heading 1')
        h.add_run(heading)
        for text in paras:
            p = doc.add_paragraph()
            set_paragraph_format(p, first_indent=True)
            set_run_font(p.add_run(text), size=12)
    for title_text, rows in TABLES:
        add_table(doc, title_text, rows)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run('物业收费管理系统 SaaS 云端商业版 · 商业上线报告'), size=9)
    DOCX_OUT.parent.mkdir(exist_ok=True)
    doc.save(DOCX_OUT)


def p_style(name, size=10.5, leading=14, align=TA_LEFT, bold=False):
    return ParagraphStyle(name, fontName='STSong-Light', fontSize=size, leading=leading, alignment=align, spaceAfter=6)


def esc(text):
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def build_pdf():
    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    doc = SimpleDocTemplate(str(PDF_OUT), pagesize=A4, leftMargin=2.54*cm, rightMargin=2.54*cm, topMargin=2.54*cm, bottomMargin=2.54*cm, pageCompression=0)
    title_style = p_style('TitleCN', size=22, leading=25, align=TA_CENTER)
    h1 = p_style('H1CN', size=14, leading=16)
    body = p_style('BodyCN', size=10.5, leading=13)
    story = [Paragraph('SaaS 云端商业版上线总报告', title_style), Paragraph('正式 Word/PDF 交付版', p_style('Sub', 11, 14, TA_CENTER)), Spacer(1, 0.3*cm)]
    for heading, paras in SECTIONS:
        story.append(Paragraph(esc(heading), h1))
        for text in paras:
            story.append(Paragraph('&nbsp;&nbsp;&nbsp;&nbsp;' + esc(text), body))
    for title_text, rows in TABLES:
        story.append(Paragraph(esc(title_text), h1))
        data = [[Paragraph(esc(c), body) for c in row] for row in rows]
        table = Table(data, colWidths=[4*cm, 11.5*cm], repeatRows=1)
        table.setStyle(TableStyle([
            ('FONTNAME', (0,0), (-1,-1), 'STSong-Light'),
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#E8EEF5')),
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('LEFTPADDING', (0,0), (-1,-1), 6), ('RIGHTPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,0), (-1,-1), 5), ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.2*cm))
    PDF_OUT.parent.mkdir(exist_ok=True)
    doc.build(story)


def main():
    build_docx()
    build_pdf()
    print(f'wrote {DOCX_OUT}')
    print(f'wrote {PDF_OUT}')


if __name__ == '__main__':
    main()
